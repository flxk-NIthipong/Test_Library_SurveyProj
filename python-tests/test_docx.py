import docx
import re

print("="*75)
print("📝 เริ่มการทดสอบ python-docx (สกัดตาราง Word พร้อมผ่าตัดตัวเลขและหน่วย)")
print("="*75)

file_name = "../MockData/word_MockData.docx"

# =====================================================================
# 🛠️ ฟังก์ชันผ่าตัดแยก "ตัวเลข" และ "หน่วย" ด้วย Regex (ยกมาจาก Pandas)
# =====================================================================
def extract_value_and_unit(text):
    # ถ้าไม่มีข้อมูลเลย ให้คืนค่าเป็น None ทั้งตัวเลขและหน่วย
    if text is None or str(text).strip() == "":
        return None, None
    
    text = str(text).strip()
    match = re.match(r'([\d\.,]+)\s*(.*)', text)
    
    if match:
        num_str = match.group(1).replace(',', '')
        try:
            num = float(num_str)
            if num.is_integer():
                num = int(num)
        except ValueError:
            num = None
            
        unit = match.group(2).strip()
        if not unit:
            unit = None
            
        return num, unit
    else:
        return None, None

# =====================================================================
# 🎯 ส่วนหลัก: อ่านไฟล์ Word และสกัดข้อมูล
# =====================================================================
try:
    print(f"📥 กำลังโหลดข้อมูลจากไฟล์: '{file_name}'...\n")
    doc = docx.Document(file_name)

    # --- ส่วนที่ 1: สกัดข้อความทั่วไป ---
    print("📄 ส่วนที่ 1: ข้อความที่อยู่นอกตาราง")
    has_text = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            print(f"   - {text}")
            has_text = True
            
    if not has_text:
        print("   (ไม่พบข้อความทั่วไปในเอกสารนี้)")

    # --- ส่วนที่ 2: สกัดและคลีนข้อมูลจากตาราง ---
    print("\n📊 ส่วนที่ 2: ข้อมูลที่อยู่ในตาราง ")
    
    if len(doc.tables) > 0:
        for table_index, table in enumerate(doc.tables):
            print(f"   === 🗂️ ตารางที่ {table_index + 1} ===")
            
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            table_data = []
            
            for row in table.rows[1:]:
                row_data = [cell.text.strip() for cell in row.cells]
                
                # ข้ามบรรทัดที่ว่างเปล่าทั้งบรรทัด
                if all(data == "" for data in row_data):
                    continue
                
                # ปั้นข้อมูลดิบเป็น Dictionary และคลีนช่องว่างเป็น None
                row_dict = {}
                for i in range(len(headers)):
                    key = headers[i] if i < len(headers) and headers[i] else f"Column_{i+1}"
                    val = row_data[i] if i < len(row_data) and row_data[i] else None
                    row_dict[key] = val

                # ผ่าตัดแยก จำนวน และ หน่วยนับ
                if 'จำนวน' in row_dict:
                    qty_val, qty_unit = extract_value_and_unit(row_dict.get('จำนวน'))
                    row_dict['จำนวน'] = qty_val
                    row_dict['หน่วยนับ'] = qty_unit

                # ผ่าตัดแยก ราคา และ หน่วยเงิน 
                # (เช็คเผื่อหัวตารางพิมพ์ว่า "ราคา" หรือ "ราคาประมาณ")
                price_key = 'ราคา' if 'ราคา' in row_dict else ('ราคาประมาณ' if 'ราคาประมาณ' in row_dict else None)
                
                if price_key:
                    price_val, price_unit = extract_value_and_unit(row_dict.get(price_key))
                    row_dict['ราคา'] = price_val
                    row_dict['หน่วยเงิน'] = price_unit
                    
                    # ถ้าระบบใช้คำว่า 'ราคาประมาณ' ให้ลบทิ้ง แล้วใช้คำว่า 'ราคา' แทนเพื่อมาตรฐานเดียวกัน
                    if price_key != 'ราคา':
                        del row_dict[price_key]

                # จัดเรียง Key ให้สวยงาม (ตัวเลือกเสริมเพื่อให้ JSON อ่านง่าย)
                ordered_dict = {}
                # ใส่ลำดับและชื่อวัสดุก่อน
                if 'ลำดับ' in row_dict: ordered_dict['ลำดับ'] = row_dict['ลำดับ']
                if 'ชื่อวัสดุ' in row_dict: ordered_dict['ชื่อวัสดุ'] = row_dict['ชื่อวัสดุ']
                # ตามด้วยจำนวนและราคา
                ordered_dict['จำนวน'] = row_dict.get('จำนวน')
                ordered_dict['หน่วยนับ'] = row_dict.get('หน่วยนับ')
                ordered_dict['ราคา'] = row_dict.get('ราคา')
                ordered_dict['หน่วยเงิน'] = row_dict.get('หน่วยเงิน')
                
                # กวาดข้อมูลที่เหลือ (ถ้ามี) มาต่อท้าย
                for k, v in row_dict.items():
                    if k not in ordered_dict:
                        ordered_dict[k] = v

                table_data.append(ordered_dict)

            # แสดงผล JSON
            print(f"   📦 ข้อมูล JSON จากตารางที่ {table_index + 1} :")
            for item in table_data:
                print(f"      {item}")
            print()
            
    else:
        print("   ❌ ไม่พบตารางใดๆ ในเอกสารนี้")

    print("✅ เสร็จสมบูรณ์!")

except FileNotFoundError:
    print(f"❌ Error: หาไฟล์ '{file_name}' ไม่เจอ")
except Exception as e:
    print(f"❌ Error: เกิดข้อผิดพลาด -> {e}")

print("="*75)