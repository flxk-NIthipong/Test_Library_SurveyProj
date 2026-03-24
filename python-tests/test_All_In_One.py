import os
import io
import json
import re 
import pandas as pd
import docx
import google.generativeai as genai
import PIL.Image
from thefuzz import process
from fastapi import FastAPI, File, UploadFile, Request, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from linebot import LineBotApi, WebhookHandler
from linebot.exceptions import InvalidSignatureError
from linebot.models import MessageEvent, TextMessage, ImageMessage, FileMessage, TextSendMessage, FlexSendMessage
import uuid
from fpdf import FPDF
from fastapi.responses import StreamingResponse, FileResponse
import tempfile
from dotenv import load_dotenv 
import google.generativeai as genai

# =================================================================
# 🔑 1. ตั้งค่า Gemini API
# =================================================================
load_dotenv()
genai.configure(api_key=os.environ.get("GEMINI_API_KEY")) # หรือใส่ API Key 

LINE_CHANNEL_ACCESS_TOKEN = os.environ.get("LINE_CHANNEL_ACCESS_TOKEN")
LINE_CHANNEL_SECRET = os.environ.get("LINE_CHANNEL_SECRET")
NGROK_URL = os.environ.get("NGROK_URL")

line_bot_api = LineBotApi(LINE_CHANNEL_ACCESS_TOKEN)
handler = WebhookHandler(LINE_CHANNEL_SECRET)
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 🗂️ ที่เก็บข้อมูลชั่วคราวเพื่อให้ React มารับไปแสดงผล
temp_data_store = {}

@app.get("/api/get-job/{job_id}")
async def get_job_data(job_id: str):
    if job_id in temp_data_store:
        return {"status": "success", "data": temp_data_store[job_id]}
    return {"status": "error", "message": "ไม่พบข้อมูล หรือข้อมูลหมดอายุแล้ว"}

# =================================================================
# 🗄️ 2. ฐานข้อมูลวัสดุมาตรฐาน (แก้คำผิด)
# =================================================================
STANDARD_MATERIALS = [
    "สายไฟ VAF 2x2.5", "สกรูเกลียวปล่อย หัว F เบอร์ 7 ยาว 2 นิ้ว", "ท่อ PVC 2 นิ้ว ชั้น 8.5",
    "เทปพันสายไฟ 3M", "สีทาบ้าน สีขาว", "หลอดไฟ LED 18W", "ปูนซีเมนต์ผสม ตราเสือ",
    "กระดาษทรายขัดไม้ เบอร์ 1", "ถุงมือผ้าเคลือบยาง", "ค้อนหงอนด้ามไฟเบอร์", "น็อตตัวเมีย M8",
    "ซิลิโคนใส กันน้ำ", "ใบตัดเหล็ก 4 นิ้ว", "สายยางรดน้ำ 5 หุน", "กาวลาเท็กซ์ TOA", "แปรงทาสี 2 นิ้ว"
]

# =================================================================
# 📥 3. API สำหรับรับไฟล์และสกัดข้อมูล
# =================================================================
@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    try:
        filename = file.filename.lower()
        content = await file.read()
        extracted_data = []

        print(f"\n📥 ได้รับไฟล์: {filename}")

        # ---------------------------------------------------------
        # 🟢 กรณีที่ 1: ไฟล์ Excel
        # ---------------------------------------------------------
        if filename.endswith(('.xlsx', '.xls')):
            print("⚙️ กำลังประมวลผลด้วย Pandas...")
            df = pd.read_excel(io.BytesIO(content))
            df = df.astype(object).where(pd.notnull(df), None) 
            raw_data = df.to_dict(orient="records")
            
            # 🌟 แปลงโครงสร้างและดึงคอลัมน์ "ราคาประมาณ" 
            for row in raw_data:
                item = {
                    "ลำดับ": row.get("ลำดับ"),
                    "ชื่อวัสดุ": row.get("ชื่อวัสดุ"),
                    "จำนวน": row.get("จำนวน"), 
                    "หน่วยนับ": None,
                    "ราคา": row.get("ราคาประมาณ") or row.get("ราคา"), 
                    "หน่วยเงิน": None
                }
                extracted_data.append(item)

        # ---------------------------------------------------------
        # 🔵 กรณีที่ 2: ไฟล์ Word
        # ---------------------------------------------------------
        elif filename.endswith('.docx'):
            print("⚙️ กำลังประมวลผลด้วย python-docx...")
            doc = docx.Document(io.BytesIO(content))
            
            if doc.tables:
                table = doc.tables[0]
                num_cols = len(table.rows[0].cells)
                for row in table.rows[1:]: 
                    cells = [cell.text.strip() for cell in row.cells]
                    if num_cols == 5:
                        row_data = {
                            "ลำดับ": cells[0], "ชื่อวัสดุ": cells[1], "จำนวน": cells[2],
                            "หน่วยนับ": None, "ราคา": cells[3], "หน่วยเงิน": None
                        }
                    else:
                        keys = ['ลำดับ', 'ชื่อวัสดุ', 'จำนวน', 'หน่วยนับ', 'ราคา', 'หน่วยเงิน']
                        row_data = {keys[i]: (cells[i] if i < len(cells) else None) for i in range(len(keys))}
                    extracted_data.append(row_data)

        # ---------------------------------------------------------
        # 🟠 กรณีที่ 3: รูปภาพ (Gemini)
        # ---------------------------------------------------------
        elif filename.endswith(('.png', '.jpg', '.jpeg')):
            print("⚙️ กำลังส่งรูปภาพให้ Gemini API อ่าน...")
            image = PIL.Image.open(io.BytesIO(content))
            model = genai.GenerativeModel('gemini-2.5-flash')
            
            prompt = """
            คุณคือ AI ผู้เชี่ยวชาญด้านการอ่านบิลและใบเสนอราคา
            กรุณาอ่านข้อมูลจากรูปภาพนี้ แล้วสกัดข้อมูลวัสดุทั้งหมดออกมาในรูปแบบ JSON Array
            โดยแต่ละรายการต้องมี Key ภาษาไทยดังนี้: "ลำดับ", "ชื่อวัสดุ", "จำนวน", "หน่วยนับ", "ราคา", "หน่วยเงิน"
            (ถ้าช่องไหนไม่มีข้อมูล ให้ใส่ค่าเป็น null)
            ส่งกลับมาแค่โครงสร้าง JSON เท่านั้น ห้ามพิมพ์คำอธิบายอื่นเพิ่มเติมเด็ดขาด
            """
            response = model.generate_content([prompt, image])
            json_str = response.text.replace("```json", "").replace("```", "").strip()
            extracted_data = json.loads(json_str)

        else:
            return {"status": "error", "message": "รองรับเฉพาะไฟล์ Excel, Word และ รูปภาพ เท่านั้น"}

        # =================================================================
        # 🌟 4. ทำความสะอาดข้อมูล: แยกจำนวน/หน่วยนับ, แยกราคา/หน่วยเงิน และแก้คำผิด
        # =================================================================
        print("🔍 กำลังทำความสะอาดข้อมูล...")
        
        for item in extracted_data:
            # --- 1. แยกจำนวนกับหน่วยนับ ---
            raw_qty = str(item.get("จำนวน") or "").strip()
            if raw_qty and raw_qty.lower() not in ["none", "nan", "-"]:
                match = re.match(r"^([\d\.]+)\s*(.*)$", raw_qty)
                if match:
                    item["จำนวน"] = match.group(1).strip()
                    if match.group(2) and not item.get("หน่วยนับ"):
                        item["หน่วยนับ"] = match.group(2).strip()

            # --- 2. 🌟 แยกราคาและหน่วยเงิน (สับตัวเลขออกจากข้อความ) ---
            raw_price = str(item.get("ราคา") or "").strip()
            raw_price = raw_price.replace(",", "") # ลบคอมมาออก (เช่น "1,200 บ." -> "1200 บ.")
            
            if raw_price and raw_price.lower() not in ["none", "nan", "-"]:
                match_price = re.match(r"^([\d\.]+)\s*(.*)$", raw_price)
                if match_price:
                    item["ราคา"] = float(match_price.group(1)) # แปลงเป็นตัวเลขเพียวๆ เพื่อให้ PDF เอาไปคำนวณได้
                    
                    # จัดระเบียบหน่วยเงินให้สวยงาม
                    if match_price.group(2):
                        currency = match_price.group(2).strip()
                        if currency in ["บ.", "THB", "฿"]: # แปลงคำแปลกๆ ให้เป็น "บาท" ให้หมด
                            currency = "บาท"
                        item["หน่วยเงิน"] = currency
                    else:
                        item["หน่วยเงิน"] = "บาท" # ถ้ามีแต่เลข ไม่มีหน่วย ให้ใส่ "บาท" อัตโนมัติ

            # --- 3. TheFuzz แก้คำผิดของชื่อวัสดุ ---
            original_name = str(item.get("ชื่อวัสดุ", "")) or ""
            item["suggestions"] = [] # เตรียมตัวเลือกไว้ให้หน้าบ้าน
            
            if original_name.strip():
                # ดึง 3 อันดับแรกที่คะแนนเกิน 60%
                matches = process.extract(original_name, STANDARD_MATERIALS, limit=3)
                suggestions = [m[0] for m in matches if m[1] >= 60]
                item["suggestions"] = suggestions
                
                best_match, score = matches[0]
                if score == 100:
                    item["สถานะ"] = "correct"
                    item["หมายเหตุ"] = "✅ ถูกต้อง"
                elif score >= 70:
                    item["สถานะ"] = "warning"
                    item["หมายเหตุ"] = f"❓ คล้ายกับ '{best_match}'"
                else:
                    item["สถานะ"] = "error"
                    item["หมายเหตุ"] = "❌ ไม่พบในฐานข้อมูล"

        print("✅ ประมวลผลสำเร็จ ส่ง JSON กลับไปที่หน้าเว็บ!")
        return {"status": "success", "data": extracted_data}

    except Exception as e:
        print("❌ เกิดข้อผิดพลาด:", str(e))
        return {"status": "error", "message": str(e)}


# =================================================================
# 📱 5. API ใหม่: สำหรับรับข้อมูลจาก LINE (Webhook)
# =================================================================
@app.post("/webhook")
async def line_webhook(request: Request):
    signature = request.headers.get("X-Line-Signature")
    body = await request.body()
    try:
        handler.handle(body.decode("utf-8"), signature)
    except InvalidSignatureError:
        raise HTTPException(status_code=400, detail="Invalid signature")
    return "OK"

# 🎯 5.1 เมื่อมีคนส่ง "ข้อความ" ทักทาย
@handler.add(MessageEvent, message=TextMessage)
def handle_text_message(event):
    reply_text = "สวัสดีครับ 📝 ส่งรูปถ่ายบิล, ไฟล์ Excel หรือ Word เข้ามาให้ผมช่วยสรุปข้อมูลจัดซื้อวัสดุได้เลยครับ!"
    line_bot_api.reply_message(event.reply_token, TextSendMessage(text=reply_text))

# 🎯 5.2 เมื่อมีคนส่ง "รูปภาพ"
@handler.add(MessageEvent, message=ImageMessage)
def handle_image_message(event):
    message_id = event.message.id
    reply_token = event.reply_token
    
    # 1. โหลดรูปจากเซิร์ฟเวอร์ LINE
    message_content = line_bot_api.get_message_content(message_id)
    image_bytes = b"".join(message_content.iter_content())
    
    try:
        # 2. นำรูปโยนเข้า Gemini
        image = PIL.Image.open(io.BytesIO(image_bytes))
        model = genai.GenerativeModel('gemini-2.5-flash')
        prompt = """
        คุณคือ AI ผู้เชี่ยวชาญด้านการอ่านบิลและใบเสนอราคา
        กรุณาอ่านข้อมูลจากรูปภาพนี้ แล้วสกัดข้อมูลวัสดุทั้งหมดออกมาในรูปแบบ JSON Array
        โดยแต่ละรายการต้องมี Key ภาษาไทยดังนี้: "ลำดับ", "ชื่อวัสดุ", "จำนวน", "หน่วยนับ", "ราคา", "หน่วยเงิน"
        (ถ้าช่องไหนไม่มีข้อมูล ให้ใส่ค่าเป็น null)
        ส่งกลับมาแค่โครงสร้าง JSON เท่านั้น ห้ามพิมพ์คำอธิบายอื่นเพิ่มเติมเด็ดขาด
        """
        response = model.generate_content([prompt, image])
        
        # 3. แปลงข้อความจาก Gemini เป็นข้อมูล JSON
        json_str = response.text.replace("```json", "").replace("```", "").strip()
        raw_data = json.loads(json_str)
        
        extracted_data = []
        for item in raw_data:
            # ทำความสะอาดข้อมูลแยกหน่วย 
            raw_qty = str(item.get("จำนวน") or "").strip()
            if raw_qty and raw_qty.lower() not in ["none", "nan", "-"]:
                match = re.match(r"^([\d\.]+)\s*(.*)$", raw_qty.replace(",", ""))
                if match:
                    item["จำนวน"] = match.group(1).strip()
                    if match.group(2) and not item.get("หน่วยนับ"):
                        item["หน่วยนับ"] = match.group(2).strip()

            raw_price = str(item.get("ราคา") or "").strip().replace(",", "") 
            if raw_price and raw_price.lower() not in ["none", "nan", "-"]:
                match_price = re.match(r"^([\d\.]+)\s*(.*)$", raw_price)
                if match_price:
                    item["ราคา"] = float(match_price.group(1))
                    if match_price.group(2):
                        currency = match_price.group(2).strip()
                        if currency in ["บ.", "THB", "฿"]: currency = "บาท"
                        item["หน่วยเงิน"] = currency
                    else:
                        item["หน่วยเงิน"] = "บาท" 
            
            extracted_data.append(item)

        # 4. สร้าง Job ID และบันทึกข้อมูล
        job_id = str(uuid.uuid4())
        temp_data_store[job_id] = extracted_data
        
        # 5. สร้าง Flex Message 
        items = [str(x.get("ชื่อวัสดุ")) for x in extracted_data if x.get("ชื่อวัสดุ")]
        total_items = len(items)
        display_limit = 5
        preview_items = items[:display_limit]
        
        body_contents = [
            {"type": "text", "text": f"📸 สกัดข้อมูลจากรูปสำเร็จ {total_items} รายการ", "weight": "bold", "size": "md", "color": "#111111", "wrap": True},
            {"type": "separator", "margin": "md"}
        ]
        
        for i, item in enumerate(preview_items):
            margin_val = "md" if i == 0 else "sm" 
            body_contents.append({"type": "text", "text": f"• {item}", "size": "sm", "color": "#333333", "wrap": True, "margin": margin_val})
        
        if total_items > display_limit:
            body_contents.append({"type": "text", "text": f"... และอื่นๆ อีก {total_items - display_limit} รายการ", "size": "xs", "color": "#999999", "margin": "md"})

        flex_message = {
            "type": "bubble",
            "header": {
                "type": "box", "layout": "vertical", "backgroundColor": "#F4F6F8",
                "contents": [{"type": "text", "text": "📄 สรุปรายการจัดซื้อวัสดุ", "weight": "bold", "size": "lg", "color": "#2C3E50"}]
            },
            "body": { "type": "box", "layout": "vertical", "spacing": "sm", "contents": body_contents },
            "footer": {
                "type": "box", "layout": "vertical", "spacing": "sm",
                "contents": [
                    {
                        "type": "button", "style": "primary", "color": "#4b49ac", 
                        "action": { "type": "uri", "label": "เปิดดูและแก้ไขในเว็บ", "uri": f"http://localhost:5173?job_id={job_id}" }
                    },
                    {
                        "type": "button", "style": "secondary", "margin": "sm",
                        "action": { "type": "uri", "label": "📥 ดาวน์โหลด PDF", "uri": f"{NGROK_URL}/api/download-pdf/{job_id}" }
                    }
                ]
            }
        }
        
        line_bot_api.reply_message(reply_token, FlexSendMessage(alt_text="สรุปรายการวัสดุจากรูปภาพ", contents=flex_message))

    except Exception as e:
        line_bot_api.reply_message(reply_token, TextSendMessage(text=f"❌ เกิดข้อผิดพลาดในการอ่านรูปภาพ: {str(e)}"))

# 🎯 5.3 เมื่อมีคนส่ง "ไฟล์ Excel หรือ Word"
@handler.add(MessageEvent, message=FileMessage)
def handle_file_message(event):
    message_id = event.message.id
    file_name = event.message.file_name.lower()
    reply_token = event.reply_token
    
    message_content = line_bot_api.get_message_content(message_id)
    file_bytes = b"".join(message_content.iter_content())
    
    try:
        raw_data = [] 
        
        # ---------------------------------------------------------
        # 🟢 กรณีที่ 1: ไฟล์ Excel
        # ---------------------------------------------------------
        if file_name.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(io.BytesIO(file_bytes))
            df = df.astype(object).where(pd.notnull(df), None) 
            raw_data = df.to_dict(orient="records")

        # ---------------------------------------------------------
        # 🔵 กรณีที่ 2: ไฟล์ Word
        # ---------------------------------------------------------
        elif file_name.endswith('.docx'):
            doc = docx.Document(io.BytesIO(file_bytes))
            if doc.tables:
                table = doc.tables[0]
                num_cols = len(table.rows[0].cells)
                for row in table.rows[1:]: 
                    cells = [cell.text.strip() for cell in row.cells]
                    if num_cols == 4:
                        raw_data.append({
                            "ลำดับ": cells[0] if len(cells) > 0 else None,
                            "ชื่อวัสดุ": cells[1] if len(cells) > 1 else None,
                            "จำนวน": cells[2] if len(cells) > 2 else None,
                            "ราคา": cells[3] if len(cells) > 3 else None
                        })
                    elif num_cols == 5:
                        raw_data.append({
                            "ลำดับ": cells[0] if len(cells) > 0 else None,
                            "ชื่อวัสดุ": cells[1] if len(cells) > 1 else None,
                            "จำนวน": cells[2] if len(cells) > 2 else None,
                            "ราคา": cells[3] if len(cells) > 3 else None
                        })
                    else:
                        keys = ['ลำดับ', 'ชื่อวัสดุ', 'จำนวน', 'หน่วยนับ', 'ราคา', 'หน่วยเงิน']
                        row_data_dict = {keys[i]: (cells[i] if i < len(cells) else None) for i in range(len(keys))}
                        raw_data.append(row_data_dict)
        else:
            line_bot_api.reply_message(reply_token, TextSendMessage(text="❌ บอทรองรับเฉพาะไฟล์ Excel, Word (.docx) และรูปภาพเท่านั้นครับ"))
            return

        # =================================================================
        # 🌟 เตรียมข้อมูลextracted_data
        # =================================================================
        extracted_data = []
        for row in raw_data:
            item = {
                "ลำดับ": row.get("ลำดับ"), 
                "ชื่อวัสดุ": row.get("ชื่อวัสดุ"),
                "จำนวน": row.get("จำนวน"), 
                "หน่วยนับ": row.get("หน่วยนับ"), 
                "ราคา": row.get("ราคาประมาณ") or row.get("ราคา"), 
                "หน่วยเงิน": row.get("หน่วยเงิน")
            }
            extracted_data.append(item)

        # =================================================================
        # 🌟 ทำความสะอาดข้อมูล: แยกจำนวน/หน่วยนับ, ราคา/หน่วยเงิน และ TheFuzz
        # =================================================================
        for item in extracted_data:
            # แยกจำนวนกับหน่วยนับ
            raw_qty = str(item.get("จำนวน") or "").strip()
            if raw_qty and raw_qty.lower() not in ["none", "nan", "-"]:
                match = re.match(r"^([\d\.]+)\s*(.*)$", raw_qty.replace(",", ""))
                if match:
                    item["จำนวน"] = match.group(1).strip()
                    if match.group(2) and not item.get("หน่วยนับ"):
                        item["หน่วยนับ"] = match.group(2).strip()

            # แยกราคาและหน่วยเงิน
            raw_price = str(item.get("ราคา") or "").strip().replace(",", "") 
            if raw_price and raw_price.lower() not in ["none", "nan", "-"]:
                match_price = re.match(r"^([\d\.]+)\s*(.*)$", raw_price)
                if match_price:
                    item["ราคา"] = float(match_price.group(1))
                    if match_price.group(2):
                        currency = match_price.group(2).strip()
                        if currency in ["บ.", "THB", "฿"]: currency = "บาท"
                        item["หน่วยเงิน"] = currency
                    else:
                        item["หน่วยเงิน"] = "บาท" 

            # TheFuzz แก้คำผิด
            original_name = str(item.get("ชื่อวัสดุ", "")) or ""
            item["suggestions"] = [] 
            if original_name.strip():
                matches = process.extract(original_name, STANDARD_MATERIALS, limit=3)
                suggestions = [m[0] for m in matches if m[1] >= 60]
                item["suggestions"] = suggestions
                
                best_match, score = matches[0]
                if score == 100:
                    item["สถานะ"] = "correct"
                elif score >= 70:
                    item["สถานะ"] = "warning"
                else:
                    item["สถานะ"] = "error"

        # --- 3. สร้างรหัส Job ID และบันทึกข้อมูลเก็บไว้ ---
        job_id = str(uuid.uuid4())
        temp_data_store[job_id] = extracted_data
        
        # --- 4. สร้าง Flex Message ---
        items = [str(x.get("ชื่อวัสดุ")) for x in extracted_data if x.get("ชื่อวัสดุ")]
        total_items = len(items)
        display_limit = 5
        preview_items = items[:display_limit]
        
        body_contents = [
            {"type": "text", "text": f"✅ สกัดข้อมูลสำเร็จ {total_items} รายการ", "weight": "bold", "size": "md", "color": "#111111", "wrap": True},
            {"type": "separator", "margin": "md"}
        ]
        
        for i, item in enumerate(preview_items):
            margin_val = "md" if i == 0 else "sm" 
            body_contents.append({"type": "text", "text": f"• {item}", "size": "sm", "color": "#333333", "wrap": True, "margin": margin_val})
        
        if total_items > display_limit:
            body_contents.append({"type": "text", "text": f"... และอื่นๆ อีก {total_items - display_limit} รายการ", "size": "xs", "color": "#999999", "margin": "md"})

        flex_message = {
            "type": "bubble",
            "header": {
                "type": "box", "layout": "vertical", "backgroundColor": "#F4F6F8",
                "contents": [{"type": "text", "text": "📄 สรุปรายการจัดซื้อวัสดุ", "weight": "bold", "size": "lg", "color": "#2C3E50"}]
            },
            "body": { "type": "box", "layout": "vertical", "spacing": "sm", "contents": body_contents },
            "footer": {
                "type": "box", "layout": "vertical", "spacing": "sm",
                "contents": [
                    {
                        "type": "button", "style": "primary", "color": "#4b49ac", 
                        "action": { "type": "uri", "label": "เปิดดูและแก้ไขในเว็บ", "uri": f"http://localhost:5173?job_id={job_id}" }
                    },
                    {
                        "type": "button", "style": "secondary", "margin": "sm",
                        "action": { "type": "uri", "label": "📥 ดาวน์โหลด PDF", "uri": f"{NGROK_URL}/api/download-pdf/{job_id}" }
                    }
                ]
            }
        }
        
        line_bot_api.reply_message(reply_token, FlexSendMessage(alt_text="สรุปรายการวัสดุที่จัดซื้อ", contents=flex_message))
        
    except Exception as e:
        line_bot_api.reply_message(reply_token, TextSendMessage(text=f"❌ เกิดข้อผิดพลาด: {str(e)}"))

# =================================================================
# 🖨️ คลาสสร้างหน้าตา PDF
# =================================================================
class ThaiPDF(FPDF):
    def __init__(self):
        super().__init__()
        try:
            self.add_font("THSarabun", "", "THSarabunNew.ttf")
            self.add_font("THSarabun", "B", "THSarabunNew Bold.ttf")
        except Exception as e:
            print("⚠️ ไม่พบไฟล์ฟอนต์ THSarabunNew.ttf หรือ THSarabunNew Bold.ttf")

    def create_material_form(self, data):
        self.add_page()
        
        # หัวกระดาษ บันทึกข้อความ
        self.set_font("THSarabun", "B", 26)
        self.cell(0, 15, "บันทึกข้อความ", ln=True, align='C')
        
        self.set_font("THSarabun", "B", 16)
        self.cell(25, 10, "ส่วนราชการ")
        self.set_font("THSarabun", "", 16)
        self.cell(0, 10, " ...................................................................................................................", ln=True)
        self.ln(5)
        
        # 📊 ตั้งค่าตาราง
        self.set_font("THSarabun", "B", 14)
        col_widths = [12, 75, 20, 25, 28, 30] 
        headers = ["ลำดับ", "รายการ", "จำนวน", "หน่วย", "ราคา/หน่วย", "รวมเงิน"]
        
        self.set_fill_color(240, 240, 240)
        for i, h in enumerate(headers):
            self.cell(col_widths[i], 10, h, border=1, align="C", fill=True)
        self.ln()

        self.set_font("THSarabun", "", 14)
        grand_total = 0
        for i, item in enumerate(data):
            name = str(item.get("ชื่อวัสดุ") or "-")
            unit = str(item.get("หน่วยนับ") or "-")
            
            try:
                qty = float(item.get("จำนวน") or 0)
            except:
                qty = 0
                
            try:
                price = float(item.get("ราคา") or 0)
            except:
                price = 0
                
            total = qty * price
            grand_total += total

            self.cell(col_widths[0], 10, str(i+1), border=1, align="C")
            self.cell(col_widths[1], 10, f" {name}", border=1)
            self.cell(col_widths[2], 10, f"{int(qty) if qty.is_integer() else qty}", border=1, align="C")
            self.cell(col_widths[3], 10, unit, border=1, align="C")
            self.cell(col_widths[4], 10, f"{price:,.2f}", border=1, align="R")
            self.cell(col_widths[5], 10, f"{total:,.2f}", border=1, align="R")
            self.ln()

        self.set_font("THSarabun", "B", 14)
        self.cell(sum(col_widths[:5]), 10, "รวมเงินทั้งสิ้น (บาท)", border=1, align="R")
        self.cell(col_widths[5], 10, f"{grand_total:,.2f}", border=1, align="R")


# =================================================================
# 📥 API สำหรับสั่งดาวน์โหลดไฟล์ PDF โดยตรง
# =================================================================
@app.get("/api/download-pdf/{job_id}")
async def download_pdf(job_id: str):
    if job_id not in temp_data_store:
        return {"status": "error", "message": "ไม่พบข้อมูล หรือข้อมูลหมดอายุแล้ว"}

    data = temp_data_store[job_id]
    
    pdf = ThaiPDF()
    pdf.create_material_form(data)
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(temp_file.name)
    
    return FileResponse(
        path=temp_file.name,
        filename="Material_Request_Summary.pdf",
        media_type="application/pdf"
    )